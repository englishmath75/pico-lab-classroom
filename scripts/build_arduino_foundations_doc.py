from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "public" / "downloads" / "arduino-foundations.docx"

FONT = "Arial"
FONT_EAST_ASIA = "Malgun Gothic"
NAVY = "0F172A"
AMBER = "FBBF24"
AMBER_LIGHT = "FFF7D6"
CYAN = "06B6D4"
CYAN_LIGHT = "ECFEFF"
SLATE = "475569"
MUTED = "64748B"
LIGHT = "F8FAFC"
LINE = "DCE3EA"
WHITE = "FFFFFF"
GREEN_LIGHT = "ECFDF5"
GREEN = "047857"


def set_run_font(run, size=11, bold=False, color=NAVY, italic=False, name=FONT):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)


def set_paragraph(p, before=0, after=6, line=1.25, align=None, keep=False):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if keep:
        p.paragraph_format.keep_with_next = True


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    flag = OxmlElement("w:tblHeader")
    flag.set(qn("w:val"), "true")
    tr_pr.append(flag)


def clear_cell(cell):
    p = cell.paragraphs[0]
    for run in list(p.runs):
        p._element.remove(run._element)
    return p


def add_cell_text(cell, text, size=10.5, bold=False, color=NAVY, align=WD_ALIGN_PARAGRAPH.LEFT, font=FONT):
    p = clear_cell(cell)
    set_paragraph(p, after=0, line=1.2, align=align)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, name=font)
    return p


def add_paragraph(doc, text, size=11, bold=False, color=NAVY, before=0, after=7, line=1.25, align=None, italic=False, keep=False):
    p = doc.add_paragraph()
    set_paragraph(p, before=before, after=after, line=line, align=align, keep=keep)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    sizes = {1: 16, 2: 13, 3: 12}
    befores = {1: 18, 2: 14, 3: 10}
    afters = {1: 10, 2: 7, 3: 5}
    colors = {1: NAVY, 2: "B45309", 3: SLATE}
    set_paragraph(p, before=befores[level], after=afters[level], line=1.15, keep=True)
    run = p.add_run(text)
    set_run_font(run, size=sizes[level], bold=True, color=colors[level])
    return p


def add_note_box(doc, label, text, fill=AMBER_LIGHT, accent="B45309"):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color=fill, size=4)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = clear_cell(cell)
    set_paragraph(p, after=0, line=1.25)
    r1 = p.add_run(f"{label}  ")
    set_run_font(r1, size=10.5, bold=True, color=accent)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=NAVY)
    add_paragraph(doc, "", size=2, after=1)
    return table


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color=NAVY, size=4)
    cell = table.cell(0, 0)
    set_cell_shading(cell, NAVY)
    p = clear_cell(cell)
    set_paragraph(p, after=0, line=1.08)
    for idx, line in enumerate(code.splitlines()):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line or " ")
        set_run_font(run, size=8.4, color="CFFAFE", name="Consolas")
    add_paragraph(doc, "", size=2, after=1)
    return table


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    set_paragraph(p, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("ARDUINO → PICO LAB  ·  Arduino Uno 기초 완성 교재  ·  ")
    set_run_font(r, size=8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.36)
    section.footer_distance = Inches(0.36)
    add_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, color in ((1, 16, NAVY), (2, 13, "B45309"), (3, 12, SLATE)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)


def add_cover(doc):
    add_paragraph(doc, "ARDUINO → PICO LAB · 사전 학습", size=10.5, bold=True, color="B45309", before=28, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "아두이노 Uno 기초 완성 교재", size=29, bold=True, color=NAVY, after=8, line=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "보드 구조 · 핀 · ADC · PWM · 시리얼 통신", size=16, bold=True, color=CYAN, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(
        doc,
        "처음 배우는 학생이 보드의 구조와 입력·처리·출력 원리를 이해하고, 실습과 중간고사를 함께 준비할 수 있도록 구성한 학생용 통합 자료입니다.",
        size=11.5,
        color=SLATE,
        after=26,
        line=1.35,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    table = doc.add_table(rows=2, cols=3)
    set_table_geometry(table, [3120, 3120, 3120])
    set_table_borders(table, color="F3D48A", size=7)
    labels = [("01", "보드와 핀 구조"), ("02", "ADC · PWM"), ("03", "코드 · 평가")]
    for idx, (number, title) in enumerate(labels):
        set_cell_shading(table.cell(0, idx), NAVY)
        add_cell_text(table.cell(0, idx), number, size=14, bold=True, color=AMBER, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.cell(1, idx), AMBER_LIGHT)
        add_cell_text(table.cell(1, idx), title, size=11, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_paragraph(doc, "", size=3, after=7)
    add_note_box(
        doc,
        "학습 방법",
        "개념을 읽은 뒤 예제 코드의 출력 결과를 먼저 예상하고, Arduino IDE 시리얼 모니터에서 실제 결과를 확인하세요.",
        fill=CYAN_LIGHT,
        accent="0E7490",
    )
    add_paragraph(doc, "대상  고등학교 소프트웨어와 생활 · Arduino 기초 실습", size=9.5, color=MUTED, before=18, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "자료 범위  Uno 구조·전원·디지털 I/O·A0~A5·10비트 ADC·8비트 PWM·시리얼 통신", size=9.5, color=MUTED, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "평가 준비  개념 비교·회로 해석·코드 실행 결과·오류 진단·서술형", size=9.5, color=MUTED, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def add_bit_section(doc):
    add_heading(doc, "1. 컴퓨터는 0과 1로 정보를 표현한다", 1)
    add_paragraph(doc, "비트(bit)는 컴퓨터가 정보를 표현하는 가장 작은 단위입니다. 한 비트에는 0 또는 1 가운데 하나만 저장됩니다. 여러 비트를 묶으면 숫자, 문자, 센서값, 명령을 표현할 수 있습니다.", size=11.2)

    table = doc.add_table(rows=5, cols=4)
    set_table_geometry(table, [1800, 2160, 2160, 3240])
    set_table_borders(table)
    headers = ["문자", "10진수(DEC)", "16진수(HEX)", "2진수(BIN)"]
    for i, text in enumerate(headers):
        set_cell_shading(table.cell(0, i), NAVY)
        add_cell_text(table.cell(0, i), text, size=10, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [("A", "65", "41", "1000001"), ("B", "66", "42", "1000010"), ("a", "97", "61", "1100001"), ("0", "48", "30", "110000")]
    for r_idx, row in enumerate(rows, 1):
        for c_idx, value in enumerate(row):
            set_cell_shading(table.cell(r_idx, c_idx), WHITE if r_idx % 2 else LIGHT)
            add_cell_text(table.cell(r_idx, c_idx), value, size=10.2, bold=c_idx == 0, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, font="Consolas" if c_idx else FONT)
    set_repeat_table_header(table.rows[0])

    add_note_box(doc, "시험 핵심", "ASCII는 문자와 숫자 코드를 대응시키는 문자 부호 체계입니다. 문자 'B'의 ASCII 10진수 값은 66이며, 이를 16진수로 나타내면 42, 2진수로 나타내면 1000010입니다.")
    add_heading(doc, "문자와 문자열의 표기", 2)
    add_paragraph(doc, "Arduino C/C++에서 작은따옴표는 문자 하나(char), 큰따옴표는 여러 문자가 이어진 문자열을 뜻합니다.")
    add_code_block(doc, "Serial.println('B');      // 문자 1개\nSerial.println(\"B\");      // 문자열\nSerial.println('B', DEC); // ASCII 값 66")
    add_note_box(doc, "주의", "화면에 보이는 글자 수와 전송 바이트 수는 항상 같지 않습니다. 영문 ASCII 문자는 보통 1바이트이지만, 한글은 UTF-8에서 보통 한 글자가 3바이트입니다.", fill=GREEN_LIGHT, accent=GREEN)
    doc.add_page_break()


def add_serial_section(doc):
    add_heading(doc, "2. Serial.begin(speed)와 통신 속도", 1)
    add_paragraph(doc, "Arduino Uno와 PC의 시리얼 모니터가 데이터를 주고받으려면 통신 속도를 먼저 정해야 합니다. setup()에서 한 번 실행하는 Serial.begin(speed)이 이 속도를 설정합니다.", size=11.2)
    add_code_block(doc, "void setup() {\n  Serial.begin(115200);\n}\n\nvoid loop() {\n}")

    table = doc.add_table(rows=4, cols=3)
    set_table_geometry(table, [2100, 2700, 4560])
    set_table_borders(table)
    headers = ["속도 예", "의미", "수업에서 확인할 점"]
    for i, text in enumerate(headers):
        set_cell_shading(table.cell(0, i), NAVY)
        add_cell_text(table.cell(0, i), text, size=10, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("9600", "초당 9,600비트", "기초 예제에서 자주 사용"),
        ("115200", "초당 115,200비트", "빠른 출력과 센서값 확인에 자주 사용"),
        ("그 밖의 값", "300, 600, 1200, 2400, 4800, 14400, 19200, 38400, 57600 등", "보드 코드와 시리얼 모니터 속도가 반드시 같아야 함"),
    ]
    for r_idx, row in enumerate(rows, 1):
        for c_idx, value in enumerate(row):
            set_cell_shading(table.cell(r_idx, c_idx), WHITE if r_idx % 2 else LIGHT)
            add_cell_text(table.cell(r_idx, c_idx), value, size=10.1, bold=c_idx == 0, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER if c_idx < 2 else WD_ALIGN_PARAGRAPH.LEFT)
    set_repeat_table_header(table.rows[0])

    add_heading(doc, "115200 bps에서 왜 약 11,520바이트인가?", 2)
    add_paragraph(doc, "일반적인 8N1 방식에서는 한 바이트를 보낼 때 시작 비트 1개, 데이터 비트 8개, 정지 비트 1개가 필요합니다. 따라서 한 바이트를 전송하는 데 모두 10비트가 사용됩니다.")
    add_note_box(doc, "계산", "115,200 bit/s ÷ 10 bit/byte = 이론상 약 11,520 byte/s", fill=CYAN_LIGHT, accent="0E7490")
    add_paragraph(doc, "이 값은 통신 방식의 오버헤드를 고려한 이론적 최대치입니다. 실제 프로그램에서는 처리 시간, 버퍼, USB 변환 등의 영향으로 체감 속도가 달라질 수 있습니다.", size=10.5, color=SLATE)
    add_note_box(doc, "자주 틀리는 표현", "115200은 '초당 115200문자'가 아니라 초당 전송하는 비트 수(bps)입니다. 영문 ASCII 1바이트 문자를 8N1로 보낼 때 이론상 약 11,520개에 해당합니다.")
    doc.add_page_break()


def add_print_section(doc):
    add_heading(doc, "3. Serial.print()와 Serial.println()", 1)
    add_paragraph(doc, "두 함수는 모두 시리얼 모니터에 값을 출력합니다. 차이는 값을 출력한 뒤 줄을 바꾸는지 여부입니다.", size=11.2)

    table = doc.add_table(rows=3, cols=3)
    set_table_geometry(table, [2700, 2700, 3960])
    set_table_borders(table)
    headers = ["함수", "출력 후 동작", "예시"]
    for i, text in enumerate(headers):
        set_cell_shading(table.cell(0, i), NAVY)
        add_cell_text(table.cell(0, i), text, size=10, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("Serial.print(val)", "같은 줄에 이어서 출력", "print(10); print(20); → 1020"),
        ("Serial.println(val)", "출력한 다음 줄바꿈", "println(10); println(20); → 10과 20이 서로 다른 줄"),
    ]
    for r_idx, row in enumerate(rows, 1):
        for c_idx, value in enumerate(row):
            set_cell_shading(table.cell(r_idx, c_idx), WHITE if r_idx % 2 else LIGHT)
            add_cell_text(table.cell(r_idx, c_idx), value, size=10.2, bold=c_idx == 0, color=NAVY, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_repeat_table_header(table.rows[0])

    add_heading(doc, "두 번째 인수 format의 의미", 2)
    add_paragraph(doc, "정수를 출력할 때는 진법, 실수를 출력할 때는 소수점 아래 자릿수를 지정합니다. 같은 숫자라도 format에 따라 화면에 보이는 표현이 달라집니다.")

    table = doc.add_table(rows=7, cols=3)
    set_table_geometry(table, [3900, 2300, 3160])
    set_table_borders(table)
    headers = ["코드", "출력", "의미"]
    for i, text in enumerate(headers):
        set_cell_shading(table.cell(0, i), NAVY)
        add_cell_text(table.cell(0, i), text, size=10, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("Serial.println(23, DEC);", "23", "10진수"),
        ("Serial.println(23, HEX);", "17", "16진수"),
        ("Serial.println(23, BIN);", "10111", "2진수"),
        ("Serial.println('B', DEC);", "66", "문자 B의 ASCII 10진수"),
        ("Serial.println('B', HEX);", "42", "문자 B의 ASCII 16진수"),
        ("Serial.println('B', BIN);", "1000010", "문자 B의 ASCII 2진수"),
    ]
    for r_idx, row in enumerate(rows, 1):
        for c_idx, value in enumerate(row):
            set_cell_shading(table.cell(r_idx, c_idx), WHITE if r_idx % 2 else LIGHT)
            add_cell_text(table.cell(r_idx, c_idx), value, size=9.8, bold=c_idx == 0, color=NAVY, align=WD_ALIGN_PARAGRAPH.LEFT if c_idx != 1 else WD_ALIGN_PARAGRAPH.CENTER, font="Consolas" if c_idx < 2 else FONT)
    set_repeat_table_header(table.rows[0])

    add_note_box(doc, "실수 출력", "Serial.println(12.235, 1)은 소수점 아래 1자리, (12.235, 2)는 2자리, (12.235, 3)은 3자리까지 표현합니다. 지정한 다음 자리의 값에 따라 반올림될 수 있습니다.", fill=GREEN_LIGHT, accent=GREEN)
    doc.add_page_break()


def add_code_reading_section(doc):
    add_heading(doc, "4. 전체 코드를 읽고 결과를 예측하기", 1)
    code = """void setup() {
  Serial.begin(115200);

  Serial.println(23, DEC);
  Serial.println(23, HEX);
  Serial.println(23, BIN);

  Serial.println('B', DEC);
  Serial.println('B', HEX);
  Serial.println('B', BIN);

  Serial.println(12.235, 1);
  Serial.println(12.235, 2);
  Serial.println(12.235, 3);
}

void loop() {
}"""
    add_code_block(doc, code)

    table = doc.add_table(rows=6, cols=3)
    set_table_geometry(table, [2600, 2980, 3780])
    set_table_borders(table)
    headers = ["코드 요소", "역할", "해석"]
    for i, text in enumerate(headers):
        set_cell_shading(table.cell(0, i), NAVY)
        add_cell_text(table.cell(0, i), text, size=10, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("void setup()", "초기 설정", "전원을 켜거나 리셋했을 때 한 번 실행"),
        ("Serial.begin(115200)", "통신 준비", "보드와 PC 사이의 통신 속도 설정"),
        ("Serial.println(23, HEX)", "정수 출력", "23을 16진수 17로 표현"),
        ("Serial.println('B', BIN)", "문자 코드 출력", "B의 ASCII 값 66을 2진수로 표현"),
        ("void loop()", "반복 실행", "비어 있으므로 반복해서 수행할 명령이 없음"),
    ]
    for r_idx, row in enumerate(rows, 1):
        for c_idx, value in enumerate(row):
            set_cell_shading(table.cell(r_idx, c_idx), WHITE if r_idx % 2 else LIGHT)
            add_cell_text(table.cell(r_idx, c_idx), value, size=10, bold=c_idx == 0, color=NAVY, font="Consolas" if c_idx == 0 else FONT)
    set_repeat_table_header(table.rows[0])

    add_heading(doc, "실습 오류를 찾는 순서", 2)
    errors = doc.add_table(rows=5, cols=2)
    set_table_geometry(errors, [2500, 6860])
    set_table_borders(errors)
    error_rows = [
        ("1. 화면이 비어 있음", "보드·포트 선택, 업로드 성공 여부, 시리얼 모니터 열기 확인"),
        ("2. 글자가 깨짐", "코드의 Serial.begin 속도와 시리얼 모니터 속도를 같게 설정"),
        ("3. 한 번만 출력됨", "출력문이 setup()에 있으면 정상이며, 반복 출력은 loop()에 작성"),
        ("4. 결과가 예상과 다름", "DEC·HEX·BIN 또는 소수점 자릿수 인수를 확인"),
        ("5. 컴파일 오류", "대소문자, 세미콜론, 괄호, 작은따옴표·큰따옴표를 점검"),
    ]
    for r_idx, row in enumerate(error_rows):
        for c_idx, value in enumerate(row):
            set_cell_shading(errors.cell(r_idx, c_idx), AMBER_LIGHT if c_idx == 0 else WHITE)
            add_cell_text(errors.cell(r_idx, c_idx), value, size=10.2, bold=c_idx == 0, color=NAVY)
    doc.add_page_break()


def add_exam_section(doc):
    add_heading(doc, "5. 실습·중간고사 대비 정리", 1)
    add_paragraph(doc, "아래 항목은 시작 전 기초에서 문제은행에 등록할 핵심 개념입니다. 이후 실습 자료가 추가되면 같은 분류에 누적하여 최종 예상문제 문서로 통합합니다.", size=11.2)

    table = doc.add_table(rows=7, cols=4)
    set_table_geometry(table, [1500, 2200, 3580, 2080])
    set_table_borders(table)
    headers = ["범위", "핵심어", "학생이 설명해야 할 내용", "출제 준비 유형"]
    for i, text in enumerate(headers):
        set_cell_shading(table.cell(0, i), NAVY)
        add_cell_text(table.cell(0, i), text, size=9.6, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("기초", "bit", "0과 1로 정보를 표현하는 최소 단위", "개념 선택"),
        ("기초", "ASCII", "문자와 숫자 코드의 대응", "변환·해석"),
        ("기초", "baud rate", "bps와 8N1의 전송량 계산", "계산·서술"),
        ("기초", "Serial.begin", "시리얼 통신 속도 설정과 위치", "빈칸·코드"),
        ("기초", "print / println", "줄바꿈 여부와 출력 결과", "출력 예측"),
        ("기초", "DEC·HEX·BIN", "정수와 문자 코드를 여러 진법으로 표현", "코드 해석"),
    ]
    for r_idx, row in enumerate(rows, 1):
        for c_idx, value in enumerate(row):
            set_cell_shading(table.cell(r_idx, c_idx), WHITE if r_idx % 2 else LIGHT)
            add_cell_text(table.cell(r_idx, c_idx), value, size=9.7, bold=c_idx in (0, 1), color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER if c_idx in (0, 1, 3) else WD_ALIGN_PARAGRAPH.LEFT)
    set_repeat_table_header(table.rows[0])

    add_heading(doc, "지금 스스로 설명해 보기", 2)
    prompts = [
        "115200 bps를 '초당 115200문자'라고 말하면 왜 정확하지 않은가?",
        "Serial.print()와 Serial.println()의 차이는 무엇인가?",
        "문자 'B'를 DEC, HEX, BIN 형식으로 출력하면 각각 무엇이 나오는가?",
        "시리얼 모니터의 글자가 깨질 때 가장 먼저 확인할 설정은 무엇인가?",
        "setup()과 loop()에 출력문을 넣었을 때 실행 횟수가 어떻게 달라지는가?",
    ]
    check = doc.add_table(rows=len(prompts), cols=2)
    set_table_geometry(check, [700, 8660])
    set_table_borders(check, color="CDE7EA")
    for idx, prompt in enumerate(prompts, 1):
        set_cell_shading(check.cell(idx - 1, 0), CYAN_LIGHT)
        add_cell_text(check.cell(idx - 1, 0), str(idx), size=11, bold=True, color="0E7490", align=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(check.cell(idx - 1, 1), prompt, size=10.3, bold=False, color=NAVY)

    add_note_box(doc, "다음 단계", "추가 캡처가 제공되면 핵심 개념, 코드 해석, 실습 결과, 오류 점검, 문제 유형을 같은 구조로 누적합니다. 전체 범위가 확정되면 객관식·서술형·코드 실행 결과 예측 문제를 한 번에 문서화합니다.", fill=AMBER_LIGHT, accent="B45309")
    add_paragraph(doc, "자료 안내  제공된 수업 캡처의 핵심 개념을 바탕으로 교실 수업 목적에 맞게 요약·재구성함.", size=9, color=MUTED, before=10, after=2)
    add_paragraph(doc, "이 문서는 화면을 복제한 자료가 아니라 개념과 코드를 독자적으로 설명한 학습자료입니다.", size=9, color=MUTED, after=2)


def add_uno_structure_section(doc):
    doc.add_page_break()
    add_heading(doc, "6. Arduino Uno 보드의 구조와 역할", 1)
    add_paragraph(doc, "Arduino Uno는 센서와 스위치의 신호를 입력받고, 프로그램에 따라 처리한 뒤 LED·부저·모터 같은 장치를 제어하는 마이크로컨트롤러 보드입니다. 수업에서는 Uno R3의 ATmega328P 기반 보드를 기준으로 설명합니다.", size=11.2)

    table = doc.add_table(rows=9, cols=3)
    set_table_geometry(table, [2350, 3500, 3510])
    set_table_borders(table)
    headers = ["보드 구성", "역할", "수업에서 확인할 점"]
    for i, text in enumerate(headers):
        set_cell_shading(table.cell(0, i), NAVY)
        add_cell_text(table.cell(0, i), text, size=10, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("ATmega328P", "작성한 스케치를 실행하는 핵심 마이크로컨트롤러", "디지털 I/O, ADC, 타이머 등을 포함"),
        ("USB 단자", "PC와 연결하여 프로그램 업로드·시리얼 통신·전원 공급", "충전 전용이 아닌 데이터 케이블 사용"),
        ("DC 전원 잭", "외부 어댑터로 보드에 전원 공급", "USB와 다른 전원 경로임"),
        ("RESET 버튼", "프로그램을 처음부터 다시 실행", "저장된 스케치는 지워지지 않음"),
        ("디지털 0~13", "HIGH/LOW 디지털 입력 또는 출력", "~ 표시 핀은 PWM 기능도 제공"),
        ("아날로그 A0~A5", "연속 전압을 ADC로 읽는 입력", "디지털 14~19로도 사용 가능"),
        ("전원 핀", "5V·3.3V·GND·VIN 등 전원 연결", "전압과 GND를 먼저 확인"),
        ("내장 LED", "D13에 연결된 보드 위 LED", "외부 부품 없이 Blink 실습 가능"),
    ]
    for r_idx, row in enumerate(rows, 1):
        for c_idx, value in enumerate(row):
            set_cell_shading(table.cell(r_idx, c_idx), WHITE if r_idx % 2 else LIGHT)
            add_cell_text(table.cell(r_idx, c_idx), value, size=9.7, bold=c_idx == 0, color=NAVY)
    set_repeat_table_header(table.rows[0])

    add_heading(doc, "입력 → 처리 → 출력으로 보는 보드", 2)
    flow = doc.add_table(rows=2, cols=3)
    set_table_geometry(flow, [3120, 3120, 3120])
    set_table_borders(flow, color="CDE7EA")
    labels = [("입력", "버튼·가변저항·센서의 값을 읽음"), ("처리", "조건·계산·반복으로 판단함"), ("출력", "LED·부저·모터를 제어함")]
    for i, (title, body) in enumerate(labels):
        set_cell_shading(flow.cell(0, i), ["E0F2FE", AMBER_LIGHT, GREEN_LIGHT][i])
        add_cell_text(flow.cell(0, i), title, size=11, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(flow.cell(1, i), body, size=9.8, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_note_box(doc, "시험 핵심", "핀 이름만 보고 입력·출력을 단정하지 않습니다. 디지털 핀은 pinMode() 설정에 따라 입력 또는 출력이 되고, A0~A5도 디지털 입출력으로 바꿔 사용할 수 있습니다.")


def add_pin_map_section(doc):
    doc.add_page_break()
    add_heading(doc, "7. Uno 핀 기능을 정확하게 구분하기", 1)
    add_paragraph(doc, "하나의 핀에 기본 기능과 특수 기능이 함께 들어 있는 경우가 많습니다. 회로를 만들기 전에는 핀 번호뿐 아니라 PWM(~), 통신, 인터럽트 같은 겸용 기능도 확인해야 합니다.", size=11.2)

    table = doc.add_table(rows=8, cols=4)
    set_table_geometry(table, [1900, 2480, 2500, 2480])
    set_table_borders(table)
    headers = ["핀 범위", "기본 역할", "대표 특수 기능", "주의"]
    for i, text in enumerate(headers):
        set_cell_shading(table.cell(0, i), NAVY)
        add_cell_text(table.cell(0, i), text, size=9.8, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("D0(RX), D1(TX)", "디지털 I/O", "USB 시리얼 통신", "업로드·통신 중 사용 충돌 주의"),
        ("D2, D3", "디지털 I/O", "외부 인터럽트", "D3은 PWM도 가능"),
        ("D3,5,6,9,10,11", "디지털 I/O", "PWM 출력(~)", "analogWrite() 사용 가능"),
        ("D10~D13", "디지털 I/O", "SPI 통신", "D13은 내장 LED와 연결"),
        ("A0~A3", "아날로그 입력", "디지털 14~17", "PWM 출력 기능은 없음"),
        ("A4(SDA)", "아날로그 입력", "디지털 18·I²C 데이터", "I²C 사용 중 겸용 주의"),
        ("A5(SCL)", "아날로그 입력", "디지털 19·I²C 클록", "I²C 사용 중 겸용 주의"),
    ]
    for r_idx, row in enumerate(rows, 1):
        for c_idx, value in enumerate(row):
            set_cell_shading(table.cell(r_idx, c_idx), WHITE if r_idx % 2 else LIGHT)
            add_cell_text(table.cell(r_idx, c_idx), value, size=9.4, bold=c_idx == 0, color=NAVY)
    set_repeat_table_header(table.rows[0])

    add_heading(doc, "전원 관련 핀", 2)
    pwr = doc.add_table(rows=7, cols=3)
    set_table_geometry(pwr, [1900, 3700, 3760])
    set_table_borders(pwr)
    headers = ["핀", "의미", "중요한 주의"]
    for i, text in enumerate(headers):
        set_cell_shading(pwr.cell(0, i), "B45309")
        add_cell_text(pwr.cell(0, i), text, size=10, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("5V", "보드의 5V 전원선", "GND와 직접 연결하면 안 됨"),
        ("3.3V", "3.3V 부품용 전원선", "허용 전류가 작으므로 큰 부하 금지"),
        ("GND", "회로 전압의 공통 기준", "외부 전원 사용 시 공통 GND 확인"),
        ("VIN", "외부 입력 전압을 보드로 공급", "5V 출력 핀과 혼동 금지"),
        ("AREF", "ADC 기준 전압 입력", "초급 실습에서는 임의 연결 금지"),
        ("IOREF/RESET", "논리 전압 기준 표시/보드 재시작", "일반 출력 핀으로 사용하지 않음"),
    ]
    for r_idx, row in enumerate(rows, 1):
        for c_idx, value in enumerate(row):
            set_cell_shading(pwr.cell(r_idx, c_idx), WHITE if r_idx % 2 else AMBER_LIGHT)
            add_cell_text(pwr.cell(r_idx, c_idx), value, size=9.7, bold=c_idx == 0, color=NAVY)
    add_note_box(doc, "안전", "Uno의 I/O 핀에는 0~5V 범위를 벗어나는 신호를 직접 넣지 않습니다. 모터처럼 전류를 많이 사용하는 부품은 GPIO에 직접 연결하지 말고 드라이버와 별도 전원을 사용하며 GND를 공통으로 연결합니다.", fill="FEE2E2", accent="991B1B")


def add_analog_input_section(doc):
    doc.add_page_break()
    add_heading(doc, "8. A0~A5: 아날로그 입력과 디지털 입출력", 1)
    add_paragraph(doc, "A0~A5의 가장 중요한 역할은 센서의 연속적인 전압을 읽는 것입니다. 그러나 이 여섯 핀은 디지털 핀 14~19로도 동작합니다. '아날로그 핀'이라는 이름 때문에 아날로그 출력도 가능하다고 오해하기 쉽지만, Uno의 A0~A5에는 DAC 방식의 아날로그 출력 기능이 없습니다.", size=11.2)

    table = doc.add_table(rows=4, cols=4)
    set_table_geometry(table, [2100, 2420, 2420, 2420])
    set_table_borders(table)
    headers = ["사용 방식", "함수", "값/상태", "예"]
    for i, text in enumerate(headers):
        set_cell_shading(table.cell(0, i), NAVY)
        add_cell_text(table.cell(0, i), text, size=10, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("아날로그 입력", "analogRead(A0)", "0~1023", "가변저항·조도 센서"),
        ("디지털 입력", "digitalRead(A0)", "LOW/HIGH", "버튼·디지털 센서"),
        ("디지털 출력", "digitalWrite(A0, ...)", "LOW/HIGH", "LED 켜기·끄기"),
    ]
    for r_idx, row in enumerate(rows, 1):
        for c_idx, value in enumerate(row):
            set_cell_shading(table.cell(r_idx, c_idx), WHITE if r_idx % 2 else LIGHT)
            add_cell_text(table.cell(r_idx, c_idx), value, size=9.7, bold=c_idx == 0, color=NAVY, font="Consolas" if c_idx == 1 else FONT)

    add_heading(doc, "10비트 ADC와 0~1023", 2)
    add_paragraph(doc, "Uno의 ADC는 기준 전압 범위를 2¹⁰=1024개의 구간으로 나눕니다. 숫자는 0부터 시작하므로 결과값은 0~1023입니다. 따라서 '1023단계'가 아니라 '1024단계, 최댓값 1023'이라고 표현해야 정확합니다.")
    add_note_box(doc, "계산", "기본 5V 기준에서 한 단계의 이상적인 크기 ≈ 5V ÷ 1024 ≈ 0.00488V = 4.88mV", fill=CYAN_LIGHT, accent="0E7490")
    add_note_box(doc, "구분", "ADC는 아날로그 전압을 디지털 숫자로 바꾸는 입력 변환기입니다. DAC는 디지털 숫자를 실제 아날로그 전압으로 바꾸는 출력 변환기인데, 일반 Uno에는 DAC 출력 핀이 없습니다.", fill=GREEN_LIGHT, accent=GREEN)

    add_heading(doc, "예제 1 · A0에서 가변저항 읽기", 2)
    add_code_block(doc, "const int SENSOR_PIN = A0;\n\nvoid setup() {\n  Serial.begin(9600);\n}\n\nvoid loop() {\n  int sensorValue = analogRead(SENSOR_PIN);\n  Serial.println(sensorValue);\n  delay(500);\n}")
    add_paragraph(doc, "가변저항 양 끝은 5V와 GND, 가운데 단자는 A0에 연결합니다. 손잡이를 돌리면 A0 전압이 변하고 sensorValue가 0~1023 범위에서 달라집니다.")
    add_heading(doc, "예제 2 · A0를 디지털 출력으로 사용", 2)
    add_code_block(doc, "void setup() {\n  pinMode(A0, OUTPUT);\n}\n\nvoid loop() {\n  digitalWrite(A0, HIGH);\n  delay(1000);\n  digitalWrite(A0, LOW);\n  delay(1000);\n}")
    add_note_box(doc, "시험 핵심", "A0라는 이름을 코드에 그대로 쓰는 편이 읽기 쉽습니다. Uno에서 A0는 디지털 14번과 같은 핀이지만, 보드 종류가 달라지면 숫자 대응이 달라질 수 있습니다.")


def add_pwm_section(doc):
    doc.add_page_break()
    add_heading(doc, "9. PWM 핀: 디지털 출력으로 밝기와 속도 조절", 1)
    add_paragraph(doc, "Uno의 D3·D5·D6·D9·D10·D11에는 ~ 표시가 있습니다. 이 핀에서는 analogWrite()로 PWM을 출력할 수 있습니다. 이름에 analog가 들어가지만, 핀이 일정한 중간 전압을 계속 내보내는 것은 아닙니다.", size=11.2)

    table = doc.add_table(rows=5, cols=4)
    set_table_geometry(table, [1700, 2100, 2600, 2960])
    set_table_borders(table)
    headers = ["analogWrite 값", "켜짐 비율", "핀의 실제 동작", "관찰 결과"]
    for i, text in enumerate(headers):
        set_cell_shading(table.cell(0, i), NAVY)
        add_cell_text(table.cell(0, i), text, size=9.7, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("0", "0%", "계속 LOW", "LED 꺼짐"),
        ("64", "약 25%", "짧게 켜지고 길게 꺼짐", "어두움"),
        ("128", "약 50%", "켜짐·꺼짐 시간이 비슷", "중간 밝기"),
        ("255", "100%", "계속 HIGH", "가장 밝음"),
    ]
    for r_idx, row in enumerate(rows, 1):
        for c_idx, value in enumerate(row):
            set_cell_shading(table.cell(r_idx, c_idx), WHITE if r_idx % 2 else LIGHT)
            add_cell_text(table.cell(r_idx, c_idx), value, size=9.7, bold=c_idx == 0, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER if c_idx < 2 else WD_ALIGN_PARAGRAPH.LEFT)

    add_heading(doc, "PWM의 핵심: 주기와 듀티비", 2)
    add_paragraph(doc, "PWM은 HIGH와 LOW를 빠르게 반복합니다. 한 주기 중 HIGH인 시간의 비율을 듀티비(duty cycle)라고 합니다. LED는 빠른 깜빡임을 평균 밝기로 느끼고, 모터는 평균적으로 전달되는 에너지의 차이로 속도가 달라집니다.")
    add_note_box(doc, "8비트 출력", "analogWrite() 값은 0~255로 모두 256단계입니다. 128은 255의 약 절반이므로 듀티비가 약 50%입니다. 이는 실제 2.5V를 계속 출력한다는 뜻이 아닙니다.", fill=CYAN_LIGHT, accent="0E7490")
    add_note_box(doc, "용어 주의", "수업에서 PWM을 편의상 '아날로그 출력'이라고 부르기도 하지만, 정확하게는 디지털 펄스의 폭을 조절해 평균 효과를 만드는 PWM 출력입니다.")

    doc.add_page_break()
    add_heading(doc, "예제 · D9 LED 밝기 서서히 바꾸기", 2)
    add_code_block(doc, "const int LED_PIN = 9;\n\nvoid setup() {\n  pinMode(LED_PIN, OUTPUT);\n}\n\nvoid loop() {\n  for (int brightness = 0; brightness <= 255; brightness++) {\n    analogWrite(LED_PIN, brightness);\n    delay(10);\n  }\n\n  for (int brightness = 255; brightness >= 0; brightness--) {\n    analogWrite(LED_PIN, brightness);\n    delay(10);\n  }\n}")
    add_paragraph(doc, "D9 → 220Ω 저항 → LED 긴 다리, LED 짧은 다리 → GND 순서로 연결합니다. brightness가 0에서 255로 증가하면 듀티비가 커져 LED가 밝아지고, 반대로 감소하면 어두워집니다.")


def add_compare_practice_section(doc):
    doc.add_page_break()
    add_heading(doc, "10. 아날로그 입력과 PWM 출력을 한 번에 연결하기", 1)
    add_paragraph(doc, "가변저항의 0~5V 전압을 A0에서 0~1023으로 읽고, 그 값을 D9의 0~255 PWM 값으로 변환해 LED 밝기를 제어합니다. 입력→처리→출력 구조를 가장 분명하게 확인할 수 있는 실습입니다.", size=11.2)
    code = '''const int SENSOR_PIN = A0;
const int LED_PIN = 9;

void setup() {
  pinMode(LED_PIN, OUTPUT);
  Serial.begin(9600);
}

void loop() {
  int sensorValue = analogRead(SENSOR_PIN);
  int pwmValue = map(sensorValue, 0, 1023, 0, 255);

  analogWrite(LED_PIN, pwmValue);
  Serial.print("ADC: ");
  Serial.print(sensorValue);
  Serial.print("  PWM: ");
  Serial.println(pwmValue);
  delay(20);
}'''
    add_code_block(doc, code)
    add_note_box(doc, "코드 해석", "analogRead()는 입력, map()은 범위 변환 처리, analogWrite()는 PWM 출력입니다. 시리얼 모니터에는 ADC 원본값과 PWM 변환값이 함께 표시됩니다.", fill=GREEN_LIGHT, accent=GREEN)

    add_heading(doc, "실습 순서", 2)
    steps = [
        ("1", "전원을 분리하고 가변저항을 5V-A0-GND에 연결한다."),
        ("2", "D9-220Ω 저항-LED-GND 순서로 출력 회로를 만든다."),
        ("3", "USB를 연결하고 보드·포트를 선택한 뒤 코드를 업로드한다."),
        ("4", "시리얼 모니터를 9600 bps로 열고 ADC와 PWM 값을 비교한다."),
        ("5", "가변저항을 돌려 입력값·변환값·LED 밝기의 관계를 기록한다."),
        ("6", "map()의 출력 범위나 delay()를 바꾸고 결과를 설명한다."),
    ]
    check = doc.add_table(rows=len(steps), cols=2)
    set_table_geometry(check, [700, 8660])
    set_table_borders(check, color="CDE7EA")
    for i, (num, text) in enumerate(steps):
        set_cell_shading(check.cell(i, 0), CYAN_LIGHT)
        add_cell_text(check.cell(i, 0), num, size=11, bold=True, color="0E7490", align=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(check.cell(i, 1), text, size=10.2, color=NAVY)

    add_heading(doc, "세 함수 비교", 2)
    table = doc.add_table(rows=4, cols=5)
    set_table_geometry(table, [1750, 2050, 1750, 1650, 2160])
    set_table_borders(table)
    headers = ["함수", "핀", "방향", "값", "정확한 의미"]
    for i, text in enumerate(headers):
        set_cell_shading(table.cell(0, i), NAVY)
        add_cell_text(table.cell(0, i), text, size=9.5, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("analogRead()", "A0~A5", "입력", "0~1023", "10비트 ADC 결과"),
        ("digitalWrite()", "디지털 I/O", "출력", "LOW/HIGH", "0V 또는 약 5V 상태"),
        ("analogWrite()", "PWM ~ 핀", "출력", "0~255", "8비트 PWM 듀티비"),
    ]
    for r_idx, row in enumerate(rows, 1):
        for c_idx, value in enumerate(row):
            set_cell_shading(table.cell(r_idx, c_idx), WHITE if r_idx % 2 else LIGHT)
            add_cell_text(table.cell(r_idx, c_idx), value, size=9.4, bold=c_idx == 0, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER if c_idx < 4 else WD_ALIGN_PARAGRAPH.LEFT, font="Consolas" if c_idx == 0 else FONT)


def add_final_assessment_section(doc):
    add_heading(doc, "11. 중간고사 대비 핵심 정리와 예상문제", 1)
    add_note_box(doc, "한 문장 정리", "Uno는 A0~A5에서 아날로그 전압을 10비트 ADC의 0~1023으로 읽고, D3·5·6·9·10·11에서는 8비트 PWM의 0~255로 출력 효과를 조절합니다.", fill=CYAN_LIGHT, accent="0E7490")

    table = doc.add_table(rows=10, cols=3)
    set_table_geometry(table, [2200, 3600, 3560])
    set_table_borders(table)
    headers = ["구분", "정확한 내용", "자주 하는 오해"]
    for i, text in enumerate(headers):
        set_cell_shading(table.cell(0, i), NAVY)
        add_cell_text(table.cell(0, i), text, size=10, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("A0~A5", "아날로그 입력 + 디지털 I/O", "아날로그 출력 핀이다"),
        ("ADC 범위", "1024단계, 결과 0~1023", "1023단계이다"),
        ("PWM 핀", "D3·5·6·9·10·11", "모든 디지털 핀에서 가능"),
        ("PWM 범위", "256단계, 값 0~255", "255단계이다"),
        ("analogWrite", "PWM 듀티비 설정", "실제 중간 전압을 계속 출력"),
        ("digitalWrite", "LOW/HIGH 두 상태", "0~255 밝기값 출력"),
        ("A4/A5", "ADC·디지털 I/O·I²C 겸용", "아날로그 입력만 가능"),
        ("D0/D1", "디지털 I/O·시리얼 RX/TX", "항상 자유롭게 사용 가능"),
        ("모터", "드라이버·외부 전원·공통 GND", "GPIO에 직접 연결"),
    ]
    for r_idx, row in enumerate(rows, 1):
        for c_idx, value in enumerate(row):
            set_cell_shading(table.cell(r_idx, c_idx), WHITE if r_idx % 2 else LIGHT)
            add_cell_text(table.cell(r_idx, c_idx), value, size=9.5, bold=c_idx == 0, color=NAVY)

    add_heading(doc, "예상문제", 2)
    questions = [
        "Uno의 A0~A5 핀이 수행할 수 있는 기능 세 가지를 쓰시오.",
        "10비트 ADC의 단계 수와 analogRead() 결과 범위를 각각 쓰시오.",
        "A0를 디지털 출력으로 설정하는 한 줄의 코드를 쓰시오.",
        "Uno에서 PWM 기능을 제공하는 디지털 핀 번호를 모두 쓰시오.",
        "analogWrite(9, 128)의 의미를 듀티비와 실제 핀 동작으로 설명하시오.",
        "PWM을 진짜 아날로그 전압 출력이라고 부르기 어려운 이유를 쓰시오.",
        "analogRead(A0)의 결과가 512일 때 5V 기준 입력 전압을 대략 계산하시오.",
        "map(sensorValue, 0, 1023, 0, 255)의 역할을 설명하시오.",
        "D0과 D1을 외부 부품에 연결할 때 특히 주의해야 하는 이유를 쓰시오.",
        "모터를 GPIO에 직접 연결하면 안 되는 이유와 안전한 연결 방법을 쓰시오.",
    ]
    qt = doc.add_table(rows=len(questions), cols=2)
    set_table_geometry(qt, [700, 8660])
    set_table_borders(qt, color="CDE7EA")
    for i, q in enumerate(questions, 1):
        set_cell_shading(qt.cell(i - 1, 0), AMBER_LIGHT)
        add_cell_text(qt.cell(i - 1, 0), str(i), size=11, bold=True, color="B45309", align=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(qt.cell(i - 1, 1), q, size=10.1, color=NAVY)

    add_heading(doc, "정답과 해설", 2)
    answers = [
        "아날로그 입력, 디지털 입력, 디지털 출력.",
        "1024단계이며 결과값은 0~1023.",
        "pinMode(A0, OUTPUT);",
        "D3, D5, D6, D9, D10, D11.",
        "D9를 약 50% 듀티비로 빠르게 HIGH/LOW 스위칭한다.",
        "중간 전압을 일정하게 출력하지 않고 0V와 5V를 빠르게 반복하기 때문이다.",
        "약 512×5/1024=2.5V.",
        "ADC의 0~1023 입력 범위를 PWM의 0~255 출력 범위로 비례 변환한다.",
        "D0·D1은 USB 시리얼 RX/TX와 겸용이어서 업로드나 통신과 충돌할 수 있다.",
        "GPIO 허용 전류보다 큰 전류가 필요하므로 모터 드라이버와 외부 전원을 사용하고 GND를 공통 연결한다.",
    ]
    at = doc.add_table(rows=len(answers), cols=2)
    set_table_geometry(at, [700, 8660])
    set_table_borders(at, color="D1FAE5")
    for i, a in enumerate(answers, 1):
        set_cell_shading(at.cell(i - 1, 0), GREEN_LIGHT)
        add_cell_text(at.cell(i - 1, 0), str(i), size=10.5, bold=True, color=GREEN, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(at.cell(i - 1, 1), a, size=9.8, color=NAVY)


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_cover(doc)
    add_bit_section(doc)
    add_serial_section(doc)
    add_print_section(doc)
    add_code_reading_section(doc)
    add_exam_section(doc)
    add_uno_structure_section(doc)
    add_pin_map_section(doc)
    add_analog_input_section(doc)
    add_pwm_section(doc)
    add_compare_practice_section(doc)
    add_final_assessment_section(doc)
    doc.core_properties.title = "아두이노 Uno 기초 완성 교재"
    doc.core_properties.subject = "Uno 구조, 핀, ADC, PWM, 시리얼 통신, 실습, 중간고사"
    doc.core_properties.author = "PICO LAB"
    doc.core_properties.keywords = "Arduino Uno, 핀, ADC, PWM, Serial, ASCII, 중간고사, 실습"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
